# dashboard/app.py
"""
Dashboard interattiva: confronto KPI andata vs ritorno (con playout per JVC)
tra le partite della stagione, sia a livello squadra che per giocatore, più un
confronto multi-giocatore per singolo KPI lungo tutta la stagione.

Avvio: streamlit run dashboard/app.py

Sezione 1 — "Confronto per KPI — Squadra / Giocatore":
per ogni avversario (13, in ordine di giornata di andata) + una 14ª posizione
per il playout (solo JVC): barre raggruppate andata/ritorno/playout-andata/
playout-ritorno, e due linee di trend sovrapposte (andata, ritorno) che
uniscono i valori delle 13 giornate di ciascun girone. Le linee sono lisciate
con un kernel gaussiano e si interrompono dove mancano dati (giocatore non in
campo); barre e/o trend sono attivabili/disattivabili.

Sezione 2 — "Confronto tra giocatori":
per un KPI scelto, una linea continua per ciascuna entità selezionata
(Squadra e/o giocatori) lungo tutte le 28 partite della stagione (andata,
ritorno, playout andata = 27ª, playout ritorno = 28ª), con tabella
riepilogativa min/mediana/max per entità.

Sezione 3 — "Confronto KPI per una singola entità":
speculare alla 2 — una entità (Squadra o un giocatore), più KPI selezionati
sovrapposti come linee continue lungo tutta la stagione. Include anche i KPI
assoluti (conteggi: punti, ace, errori per fondamentale, non solo percentuali).

KPI "Attacco SO" (side-out) e "Contrattacco": efficienza/totale/punti/errori/
murati sui due sottoinsiemi di attacchi determinati da
separate_attacks_counterattacks — SO = immediatamente dopo una ricezione con
voto scelto dalla sidebar ("Esito ricezione per Attacco SO"), Contrattacco =
tutti gli altri. Ricalcolati ad ogni cambio del filtro (vedi
src.leg_comparison.build_attacco_so_dataset), a differenza degli altri KPI
che sono fissi per la stagione.

Soglia minima attacchi (sidebar): per i KPI della famiglia "attacco"
(Attacco*, Attacco SO*, Contrattacco*), se un giocatore in una partita non
raggiunge questo numero di attacchi della stessa famiglia, per quella
partita/KPI risulta senza dati (vedi
src.leg_comparison.apply_min_attacks_threshold) — non un valore a zero.

Sezione 4 — "Pagella giocatore" (src.player_report, in corso — base dati,
non ancora una vera pagella narrativa): una tab per giocatore Decimo
riconosciuto, con presenze/assenze (src.attendance) e i finding di
rendimento più notevoli (streak di partite consecutive sopra/sotto la
mediana stagionale, o cambio di livello nel punto di rottura più marcato
della stagione) su un sottoinsieme curato di KPI — ognuno con un grafico
di evidenza (la serie grezza del KPI lungo la stagione, con la porzione
del finding evidenziata in verde/rosso). Non applica la soglia minima
attacchi della sidebar (i
finding sono calcolati sulla serie completa, i grafici devono restare
coerenti con quella stessa serie).
"""
import io
import os
import sys

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from plotly.subplots import make_subplots

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if REPO_ROOT not in sys.path:
    sys.path.insert(0, REPO_ROOT)

from src.leg_comparison import (
    ALL_KPIS,
    ATTACCO_FB_KPI_LABELS,
    ATTACCO_SO_KPI_LABELS,
    CONTRATTACCO_KPI_LABELS,
    DEFAULT_REC_VOTE,
    PERCENT_KPIS,
    apply_min_attacks_threshold,
    build_attacco_so_dataset,
    build_comparison_dataset,
    build_match_outcomes_dataset,
    get_x_axis_order,
    load_all_matches,
)
from src.player_report import build_player_report_base
from src.player_season_report import build_player_season_report
from src.setter_report import ATTACCO_ALZATO_ESCL_KPI_LABELS, ATTACCO_ALZATO_POS_KPI_LABELS, build_setter_kpi_dataset

SEASON = "2025-2026"
N_REGULAR_MATCHES = 28  # 13 andata + 13 ritorno + 2 playout (POA, POR)
REC_VOTE_OPTIONS = ["#", "+", "!", "-", "/", "="]

# ALL_KPIS + i KPI "Attacco SO"/"Attacco FB"/"Contrattacco" (parametrici
# sull'esito ricezione, vedi sidebar — Attacco FB e Contrattacco non
# dipendono direttamente dal filtro ma si spostano insieme ad Attacco SO,
# essendo complementari sullo stesso split) — lista mostrata in tutti i
# multiselect della dashboard.
DISPLAY_KPIS = (
    ALL_KPIS
    + list(ATTACCO_SO_KPI_LABELS.values())
    + list(ATTACCO_FB_KPI_LABELS.values())
    + list(CONTRATTACCO_KPI_LABELS.values())
)

# ---------------------------------------------------------------------------
# Palette — hue categoriche validate CVD-safe (vedi skill dataviz / palette.md),
# in due varianti (chiara/scura) perché il tema Streamlit attivo (light/dark)
# non deve mai lasciare testo scuro su sfondo scuro o viceversa: passiamo
# theme=None a st.plotly_chart per avere pieno controllo, e ricalcoliamo i
# colori ad ogni render in base a st.context.theme.type.
# Sezione 1: andata/ritorno restano lo stesso colore sia in barra che in linea
# di trend (il colore segue l'entità "girone"); i playout usano la stessa hue
# con opacità ridotta (encoding secondario) invece di nuovi colori.
# Sezione 2/3: colore per entità/KPI assegnato in ordine fisso sulla lista
# completa delle opzioni disponibili, non su quelle selezionate — così il
# filtro non ridipinge le serie superstiti.
# ---------------------------------------------------------------------------
PALETTE_LIGHT = ["#2a78d6", "#eb6834", "#1baf7a", "#eda100", "#e87ba4", "#008300", "#4a3aa7", "#e34948"]
PALETTE_DARK = ["#3987e5", "#d95926", "#199e70", "#c98500", "#d55181", "#008300", "#9085e9", "#e66767"]
ENTITY_DASHES = ["solid", "dash", "dot", "dashdot"]

LEG_LABELS = {"A": "Andata", "R": "Ritorno", "POA": "Playout andata", "POR": "Playout ritorno"}
LEG_ORDER = ["A", "R", "POA", "POR"]

# Palette di stato (vinta/persa) — fissa, non dipende dal tema (come da dataviz
# skill: i colori di stato non sono mai tematizzati né riusati come colori di
# serie), usata dalla striscia risultato sotto i grafici 2/3.
WIN_COLOR = "#0ca30c"
LOSS_COLOR = "#d03b3b"
UNKNOWN_COLOR = "#898781"


def theme_colors():
    """Colori/palette dipendenti dal tema Streamlit attivo (light/dark)."""
    is_dark = (st.context.theme.type or "light") == "dark"
    return {
        "surface": "#1a1a19" if is_dark else "#fcfcfb",
        "grid": "#2c2c2a" if is_dark else "#e1e0d9",
        "ink": "#ffffff" if is_dark else "#0b0b0b",
        "palette": PALETTE_DARK if is_dark else PALETTE_LIGHT,
    }


def hex_to_rgba(hex_color, alpha):
    h = hex_color.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return f"rgba({r},{g},{b},{alpha})"

# Abbreviazioni per i nomi avversario più lunghi, usate come etichette sull'asse X
# del grafico "Confronto tra giocatori" (28 posizioni, poco spazio per ciascuna).
OPPONENT_ABBR = {
    "Monterotondo": "M.tondo",
    "Stella Mantus": "S.Mantus",
    "Civitavecchia": "Civitav.",
    "Astrolabio": "Astrol.",
    "Poolstars": "Poolst.",
}


def abbr_opponent(name):
    return OPPONENT_ABBR.get(name, name)


def match_seq_tick_labels(opponent_order):
    """
    Etichette asse X del grafico "Confronto tra giocatori": nome avversario
    (abbreviato se necessario), ripetuto per andata (0..12) e ritorno (13..25);
    playout (26, 27) con suffisso per distinguere andata/ritorno.
    """
    labels = [abbr_opponent(o) for o in opponent_order]
    labels += [abbr_opponent(o) for o in opponent_order]
    labels += [f"{abbr_opponent('JVC')} (PO-A)", f"{abbr_opponent('JVC')} (PO-R)"]
    return labels

st.set_page_config(page_title="Decimo Roma — Andata vs Ritorno", layout="wide")


@st.cache_data(show_spinner="Caricamento partite della stagione da Google Drive...")
def load_matches_cached(season):
    """Cache separata dal resto: è la parte lenta (lettura Excel da Drive), riusata
    sia dal dataset KPI fisso sia da quello parametrico "Attacco SO"."""
    return load_all_matches(season)


@st.cache_data(show_spinner=False)
def load_data(season):
    matches = load_matches_cached(season)
    return build_comparison_dataset(season, matches=matches)


@st.cache_data(show_spinner=False)
def load_attacco_so_data(season, rec_vote):
    matches = load_matches_cached(season)
    return build_attacco_so_dataset(season, rec_vote=rec_vote, matches=matches)


@st.cache_data(show_spinner=False)
def load_match_outcomes(season):
    matches = load_matches_cached(season)
    return build_match_outcomes_dataset(season, matches=matches)


@st.cache_data(show_spinner=False)
def load_setter_kpi_data(season):
    """"Attacco Alzato" (ruolo palleggiatore, src.setter_report) — NON riusa
    load_matches_cached: quel `matches` è già filtrato/normalizzato, non
    adatto alla ricostruzione formazioni/identificazione palleggiatore (vedi
    docstring di src.setter_report)."""
    return build_setter_kpi_dataset(season)


@st.cache_data(show_spinner="Calcolo pagella giocatore...")
def load_player_report(season, rec_vote):
    matches = load_matches_cached(season)
    return build_player_report_base(season, rec_vote=rec_vote, matches=matches)


@st.cache_data(show_spinner="Calcolo report di sintesi stagionale...")
def load_player_season_report(season, rec_vote):
    matches = load_matches_cached(season)
    return build_player_season_report(season, rec_vote=rec_vote, matches=matches)


def kernel_smooth(y, bandwidth):
    """
    Smoothing gaussiano (Nadaraya-Watson) su una serie con eventuali NaN.
    Ogni punto valido viene ricalcolato come media pesata (kernel gaussiano
    sulla distanza in posizione) di tutti gli altri punti validi della serie;
    le posizioni originariamente NaN restano NaN — il "buco" nella linea
    (giocatore non in campo quella partita) resta visibile.
    bandwidth <= 0 disattiva lo smoothing (ritorna la serie grezza).
    """
    y = np.asarray(y, dtype=float)
    valid = ~np.isnan(y)
    if bandwidth <= 0 or valid.sum() < 2:
        return y.copy()

    idx = np.arange(len(y))
    valid_idx = idx[valid]
    valid_y = y[valid]
    out = np.full(len(y), np.nan)
    for i in valid_idx:
        d = valid_idx - i
        w = np.exp(-0.5 * (d / bandwidth) ** 2)
        out[i] = np.sum(w * valid_y) / np.sum(w)
    return out


def categorical_style(key, master_order):
    """
    Colore + tratteggio fissi per una chiave categorica (entità o KPI), in base
    alla sua posizione nell'elenco completo (non in quello filtrato/selezionato)
    — così cambiare selezione non ridipinge le serie superstiti.
    """
    palette = theme_colors()["palette"]
    i = master_order.index(key)
    color = palette[i % len(palette)]
    dash = ENTITY_DASHES[(i // len(palette)) % len(ENTITY_DASHES)]
    return color, dash


def base_layout(title, xaxis_title, yaxis_title, extra_xaxis=None):
    colors = theme_colors()
    xaxis = dict(gridcolor=colors["grid"])
    if extra_xaxis:
        xaxis.update(extra_xaxis)
    return dict(
        title=title,
        xaxis_title=xaxis_title,
        yaxis_title=yaxis_title,
        xaxis=xaxis,
        yaxis=dict(gridcolor=colors["grid"], zerolinecolor=colors["grid"]),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="left", x=0),
        plot_bgcolor=colors["surface"],
        paper_bgcolor=colors["surface"],
        font=dict(family="system-ui, -apple-system, 'Segoe UI', sans-serif", color=colors["ink"]),
        margin=dict(t=90),
        hovermode="x unified",
    )


def add_result_strip(fig, match_outcomes_df, x_positions, tick_labels, row=2, col=1):
    """
    Aggiunge, alla riga subplot indicata, una striscia di marker quadrati:
    verde = partita vinta, rosso = persa, grigio = esito indeterminato (vedi
    partita_vinta in compute_match_outcome — capita quando manca il file dei
    risultati ufficiali per quella partita) — alle stesse posizioni X
    (match_seq) del grafico principale sopra, così il confronto è immediato.
    """
    lookup = match_outcomes_df.set_index("match_seq")["partita_vinta"]
    colors, texts = [], []
    for x in x_positions:
        esito = lookup.get(x)
        # NON usare 'is True'/'is False': un valore booleano estratto da una
        # colonna pandas è un numpy.bool_, non un bool nativo, e
        # 'numpy.True_ is True' vale False in Python — con 'is' la striscia
        # risultava sempre grigia, qualunque fosse il vero esito (bug
        # riscontrato in verifica il 2026-08-12, mai notato prima perché la
        # differenza visiva grigio/colorato passava inosservata a uno sguardo
        # superficiale). 'esito is None' resta corretto per la chiave
        # mancante: Series.get() ritorna None (non NaN) di default.
        if esito is None:
            colors.append(UNKNOWN_COLOR)
            texts.append("Esito indeterminato")
        elif esito:
            colors.append(WIN_COLOR)
            texts.append("Vinta")
        else:
            colors.append(LOSS_COLOR)
            texts.append("Persa")

    fig.add_trace(
        go.Scatter(
            x=x_positions, y=[0] * len(x_positions), mode="markers",
            marker=dict(color=colors, size=14, symbol="square"),
            text=texts, hovertemplate="%{text}<extra></extra>",
            showlegend=False,
        ),
        row=row, col=col,
    )
    fig.update_yaxes(visible=False, range=[-1, 1], row=row, col=col)
    fig.update_xaxes(
        tickmode="array", tickvals=x_positions, ticktext=tick_labels,
        row=row, col=col,
    )
    fig.update_xaxes(showticklabels=False, row=row - 1, col=col)


def build_kpi_chart(df, kpi, entity_label, x_axis_order, show_bars, show_trend, bandwidth):
    """Grafico per avversario: barre raggruppate per leg e/o 2 linee di trend (andata/ritorno)."""
    sub = df[df["kpi"] == kpi]
    if sub.empty:
        return None

    palette = theme_colors()["palette"]
    color_a, color_r = palette[0], palette[1]
    leg_colors = {
        "A": color_a, "R": color_r,
        "POA": hex_to_rgba(color_a, 0.45), "POR": hex_to_rgba(color_r, 0.45),
    }

    fig = go.Figure()

    if show_bars:
        for leg in LEG_ORDER:
            leg_df = sub[sub["leg"] == leg].set_index("x_label").reindex(x_axis_order).dropna(subset=["value"])
            if leg_df.empty:
                continue
            fig.add_trace(go.Bar(
                x=leg_df.index, y=leg_df["value"],
                name=LEG_LABELS[leg], marker_color=leg_colors[leg],
                customdata=[format_kpi_value(kpi, v) for v in leg_df["value"]],
                hovertemplate="%{x}<br>%{customdata}<extra>" + LEG_LABELS[leg] + "</extra>",
            ))

    if show_trend:
        for leg, color in (("A", color_a), ("R", color_r)):
            leg_df = sub[sub["leg"] == leg].set_index("x_label").reindex(x_axis_order)
            values = leg_df["value"].to_numpy()
            if np.count_nonzero(~np.isnan(values)) < 2:
                continue
            smoothed = kernel_smooth(values, bandwidth)
            fig.add_trace(go.Scatter(
                x=x_axis_order, y=smoothed, mode="lines+markers",
                name=f"Trend {LEG_LABELS[leg].lower()}",
                line=dict(color=color, width=2, shape="spline", smoothing=0.8),
                marker=dict(size=6), connectgaps=False,
                customdata=[format_kpi_value(kpi, v) if not np.isnan(v) else "" for v in values],
                hovertemplate="%{x}<br>valore reale: %{customdata}<extra>Trend " + LEG_LABELS[leg].lower() + "</extra>",
            ))

    fig.update_layout(
        barmode="group",
        **base_layout(
            f"{kpi} — {entity_label}", "Avversario", kpi,
            extra_xaxis=dict(categoryorder="array", categoryarray=x_axis_order),
        ),
    )
    return fig


def build_player_comparison_chart(
    kpi, selected_entities, team_df, player_df, entity_master_order, bandwidth, opponent_order,
    match_outcomes_df,
):
    """
    Grafico di confronto multi-entità: una linea continua per entità su tutte
    le 28 partite, con una striscia vinta/persa sotto (vedi add_result_strip).
    """
    fig = make_subplots(rows=2, cols=1, shared_xaxes=True, row_heights=[0.85, 0.15], vertical_spacing=0.04)
    match_seq_full = list(range(N_REGULAR_MATCHES))
    stats_rows = []

    for entity in selected_entities:
        df = team_df if entity == "Squadra" else player_df[player_df["player"] == entity]
        sub = df[df["kpi"] == kpi].set_index("match_seq").reindex(match_seq_full)
        values = sub["value"].to_numpy()

        if np.count_nonzero(~np.isnan(values)) == 0:
            continue

        color, dash = categorical_style(entity, entity_master_order)
        smoothed = kernel_smooth(values, bandwidth)
        fig.add_trace(go.Scatter(
            x=match_seq_full, y=smoothed, mode="lines+markers",
            name=entity, line=dict(color=color, width=2, dash=dash, shape="spline", smoothing=0.8),
            marker=dict(size=5), connectgaps=False,
            customdata=[format_kpi_value(kpi, v) if not np.isnan(v) else "" for v in values],
            hovertemplate="%{x}<br>valore reale: %{customdata}<extra>" + entity + "</extra>",
        ), row=1, col=1)

        raw_valid = values[~np.isnan(values)]
        stats_rows.append({
            "Entità": entity,
            "Min": float(np.min(raw_valid)),
            "Mediana": float(np.median(raw_valid)),
            "Max": float(np.max(raw_valid)),
        })

    tick_labels = match_seq_tick_labels(opponent_order)
    add_result_strip(fig, match_outcomes_df, match_seq_full, tick_labels, row=2, col=1)
    fig.update_layout(**base_layout(f"Confronto giocatori — {kpi}", "Avversario", kpi))
    return fig, pd.DataFrame(stats_rows)


# PERCENT_KPIS (src.leg_comparison) copre solo i 32 KPI "fissi" della
# dashboard — "Attacco Alzato%" (ruolo palleggiatore, src.setter_report) non
# ne fa parte ma va comunque formattato come percentuale nella pagella.
PAGELLA_PERCENT_KPIS = set(PERCENT_KPIS) | {ATTACCO_ALZATO_POS_KPI_LABELS["eff"], ATTACCO_ALZATO_ESCL_KPI_LABELS["eff"]}


def format_kpi_value(kpi, value):
    """Formatta un valore GREZZO di singola partita (sempre un intero per i
    conteggi), per tabelle/hover — vedi format_kpi_average per le medie."""
    if kpi in PAGELLA_PERCENT_KPIS:
        return f"{value:.1f}%"
    return f"{value:.0f}"


def format_kpi_average(kpi, value):
    """
    Formatta una MEDIA (mediana stagionale, media di uno streak, media
    prima/seconda metà) — a differenza di format_kpi_value, i conteggi non
    sono arrotondati a intero: una media come 0.15 conteggi/partita non va
    confusa con "0" (bug riscontrato e corretto nella pagella giocatore:
    "Ricezione = (ace subiti): da 0 a 0" con format_kpi_value su medie
    frazionarie — la differenza reale (es. 0.15 vs 0.35) restava invisibile).
    """
    if kpi in PAGELLA_PERCENT_KPIS:
        return f"{value:.1f}%"
    return f"{value:.2f}"


def build_kpi_comparison_chart(
    entity, selected_kpis, team_df, player_df, kpi_master_order, bandwidth, opponent_order,
    match_outcomes_df,
):
    """
    Grafico di confronto multi-KPI per una singola entità: una linea continua
    per KPI su tutte le 28 partite, con una striscia vinta/persa sotto (vedi
    add_result_strip).
    """
    fig = make_subplots(rows=2, cols=1, shared_xaxes=True, row_heights=[0.85, 0.15], vertical_spacing=0.04)
    match_seq_full = list(range(N_REGULAR_MATCHES))
    stats_rows = []

    df = team_df if entity == "Squadra" else player_df[player_df["player"] == entity]

    for kpi in selected_kpis:
        sub = df[df["kpi"] == kpi].set_index("match_seq").reindex(match_seq_full)
        values = sub["value"].to_numpy()

        if np.count_nonzero(~np.isnan(values)) == 0:
            continue

        color, dash = categorical_style(kpi, kpi_master_order)
        smoothed = kernel_smooth(values, bandwidth)
        fig.add_trace(go.Scatter(
            x=match_seq_full, y=smoothed, mode="lines+markers",
            name=kpi, line=dict(color=color, width=2, dash=dash, shape="spline", smoothing=0.8),
            marker=dict(size=5), connectgaps=False,
            customdata=[format_kpi_value(kpi, v) if not np.isnan(v) else "" for v in values],
            hovertemplate="%{x}<br>valore reale: %{customdata}<extra>" + kpi + "</extra>",
        ), row=1, col=1)

        raw_valid = values[~np.isnan(values)]
        stats_rows.append({
            "KPI": kpi,
            "Min": format_kpi_value(kpi, np.min(raw_valid)),
            "Mediana": format_kpi_value(kpi, np.median(raw_valid)),
            "Max": format_kpi_value(kpi, np.max(raw_valid)),
        })

    tick_labels = match_seq_tick_labels(opponent_order)
    add_result_strip(fig, match_outcomes_df, match_seq_full, tick_labels, row=2, col=1)
    fig.update_layout(**base_layout(f"Confronto KPI — {entity}", "Avversario", "Valore"))
    return fig, pd.DataFrame(stats_rows)


def render_kpi_section(team_df, player_df, x_axis_order, entity_options):
    st.header("Confronto per KPI — Squadra / Giocatore")

    with st.sidebar:
        st.subheader("Confronto per KPI")
        selected_kpis = st.multiselect("KPI", DISPLAY_KPIS, default=[DISPLAY_KPIS[0]], key="kpi1")
        selected_entities = st.multiselect("Squadra / Giocatori", entity_options, default=["Squadra"], key="entity1")
        display_mode = st.radio("Mostra", ["Barre e trend", "Solo barre", "Solo trend"], key="mode1")
        show_table = st.checkbox("Mostra tabella dati sotto ogni grafico", key="table1")

    if not selected_kpis or not selected_entities:
        st.info("Seleziona almeno un KPI e un'entità (Squadra o un giocatore) dalla barra laterale.")
        return

    show_bars = display_mode != "Solo trend"
    show_trend = display_mode != "Solo barre"

    for entity in selected_entities:
        st.subheader(entity)
        entity_df = team_df if entity == "Squadra" else player_df[player_df["player"] == entity]

        cols = st.columns(2) if len(selected_kpis) > 1 else [st.container()]
        for i, kpi in enumerate(selected_kpis):
            fig = build_kpi_chart(
                entity_df, kpi, entity, x_axis_order, show_bars, show_trend,
                st.session_state["bandwidth"],
            )
            with cols[i % len(cols)]:
                if fig is None:
                    st.warning(f"Nessun dato per {kpi} — {entity}.")
                    continue
                st.plotly_chart(fig, theme=None)
                if show_table:
                    table = (
                        entity_df[entity_df["kpi"] == kpi][["x_label", "leg", "value"]]
                        .rename(columns={"x_label": "avversario"})
                        .sort_values(["avversario", "leg"])
                        .reset_index(drop=True)
                    )
                    st.dataframe(table, width="stretch")


def render_comparison_section(team_df, player_df, entity_options, opponent_order, match_outcomes_df):
    st.header("Confronto tra giocatori")
    st.caption("Linea continua per entità lungo tutta la stagione (asse X: avversario, andata poi ritorno; in coda il playout, sempre JVC) — striscia sotto: esito partita (verde=vinta, rosso=persa)")

    col_a, col_b = st.columns(2)
    with col_a:
        compare_kpis = st.multiselect("KPI", DISPLAY_KPIS, default=[DISPLAY_KPIS[0]], key="kpi2")
    with col_b:
        default_entities = entity_options[:3] if len(entity_options) >= 3 else entity_options
        compare_entities = st.multiselect(
            "Squadra / Giocatori da confrontare", entity_options, default=default_entities, key="entity2"
        )

    if not compare_kpis or not compare_entities:
        st.info("Seleziona almeno un KPI e almeno un'entità da confrontare.")
        return

    for kpi in compare_kpis:
        fig, stats_df = build_player_comparison_chart(
            kpi, compare_entities, team_df, player_df, entity_options, st.session_state["bandwidth"],
            opponent_order, match_outcomes_df,
        )
        st.plotly_chart(fig, theme=None)

        if not stats_df.empty:
            fmt = "%.1f%%" if kpi in PERCENT_KPIS else "%.0f"
            st.dataframe(
                stats_df,
                column_config={
                    "Min": st.column_config.NumberColumn(format=fmt),
                    "Mediana": st.column_config.NumberColumn(format=fmt),
                    "Max": st.column_config.NumberColumn(format=fmt),
                },
                width="stretch",
                hide_index=True,
            )
        else:
            st.warning(f"Nessun dato per {kpi} sulle entità selezionate.")


def render_kpi_comparison_section(team_df, player_df, entity_options, opponent_order, match_outcomes_df):
    st.header("Confronto KPI per una singola entità")
    st.caption("Linea continua per KPI, per Squadra o un giocatore, lungo tutta la stagione — striscia sotto: esito partita (verde=vinta, rosso=persa)")

    col_a, col_b = st.columns(2)
    with col_a:
        compare_entity = st.selectbox("Squadra / Giocatore", entity_options, key="entity3")
    with col_b:
        default_kpis = DISPLAY_KPIS[:3]
        compare_kpis2 = st.multiselect("KPI da confrontare", DISPLAY_KPIS, default=default_kpis, key="kpi3")

    if not compare_kpis2:
        st.info("Seleziona almeno un KPI da confrontare.")
        return

    if any(k in PERCENT_KPIS for k in compare_kpis2) and any(k not in PERCENT_KPIS for k in compare_kpis2):
        st.caption("⚠️ Stai confrontando KPI percentuali e conteggi assoluti sullo stesso asse — le scale non sono direttamente comparabili.")

    fig, stats_df = build_kpi_comparison_chart(
        compare_entity, compare_kpis2, team_df, player_df, DISPLAY_KPIS, st.session_state["bandwidth"],
        opponent_order, match_outcomes_df,
    )
    st.plotly_chart(fig, theme=None)

    if not stats_df.empty:
        st.dataframe(stats_df, width="stretch", hide_index=True)
    else:
        st.warning(f"Nessun dato per {compare_entity} sui KPI selezionati.")


_LEG_WORD = {"A": "andata", "R": "ritorno", "POA": "playout andata", "POR": "playout ritorno"}


def build_match_label_lookup(team_df):
    """
    match_seq -> 'Avversario (andata/ritorno/playout andata/playout ritorno)'
    — per riferirsi ai finding della pagella per squadra e leg invece che per
    numero di partita nudo (richiesto esplicitamente dall'utente il
    2026-08-13; il numero resta disponibile tra parentesi in format_finding_text).
    """
    rows = team_df[["match_seq", "opponent", "leg"]].drop_duplicates("match_seq").set_index("match_seq")
    return {
        ms: f"{abbr_opponent(row['opponent'])} ({_LEG_WORD[row['leg']]})"
        for ms, row in rows.iterrows()
    }


def _match_ref(ms, match_label_lookup):
    """'partita N' come fallback se il match_seq non è nel lookup (non
    dovrebbe succedere con dati validi, ma meglio un fallback esplicito che
    un KeyError)."""
    label = match_label_lookup.get(ms, f"partita {ms + 1}")
    return f"{label} (partita {ms + 1})"


def format_finding_text(finding, match_label_lookup):
    """Frase discorsiva per un finding di src.player_report (streak o cambio
    di livello) — vedi build_finding_chart per l'evidenza grafica."""
    kpi = finding["kpi"]
    direzione = finding["direzione"]
    if finding["tipo_finding"] == "streak":
        return (
            f"**{kpi}** — periodo **{direzione}** da {_match_ref(finding['start_match_seq'], match_label_lookup)} "
            f"a {_match_ref(finding['end_match_seq'], match_label_lookup)} ({finding['length']} partite di fila): "
            f"media {format_kpi_average(kpi, finding['mean_value'])} contro una mediana stagionale di "
            f"{format_kpi_average(kpi, finding['baseline'])}."
        )
    return (
        f"**{kpi}** — cambio di livello **{direzione}** da {_match_ref(finding['split_match_seq'], match_label_lookup)} "
        f"in poi: da {format_kpi_average(kpi, finding['prima_media'])} a "
        f"{format_kpi_average(kpi, finding['dopo_media'])} "
        f"(pendenza di regressione: {finding['slope']:+.3f} per partita)."
    )


def _add_finding_annotation(fig, finding, colors):
    """
    Aggiunge a `fig` l'evidenza grafica di UN finding (rettangolo per uno
    streak, linea di spartiacque + medie per un cambio di livello) — verde
    se favorevole/rosso se sfavorevole (stessa convenzione status color di
    add_result_strip). Fattorizzata da build_finding_chart per essere
    riusata anche da build_unified_finding_chart (più finding sullo stesso
    KPI in un unico grafico, per il report Word — vedi generate_player_report_docx).
    """
    highlight_color = WIN_COLOR if finding["direzione"] == "positivo" else LOSS_COLOR
    if finding["tipo_finding"] == "streak":
        fig.add_vrect(
            x0=finding["start_match_seq"] - 0.5, x1=finding["end_match_seq"] + 0.5,
            fillcolor=hex_to_rgba(highlight_color, 0.18), line_width=0,
        )
        fig.add_hline(
            y=finding["baseline"], line_dash="dot", line_color=colors["grid"],
            annotation_text="mediana stagionale", annotation_position="bottom right",
        )
    else:
        # split_match_seq è la PRIMA partita del "dopo" — la linea di
        # spartiacque va appena prima di quel punto.
        split = finding["split_match_seq"] - 0.5
        fig.add_vline(x=split, line_dash="dash", line_color=colors["grid"], annotation_text="cambio di livello")
        fig.add_shape(
            type="line", x0=0, x1=split,
            y0=finding["prima_media"], y1=finding["prima_media"],
            line=dict(color=colors["grid"], dash="dot", width=2),
        )
        fig.add_shape(
            type="line", x0=split, x1=N_REGULAR_MATCHES - 1,
            y0=finding["dopo_media"], y1=finding["dopo_media"],
            line=dict(color=highlight_color, dash="dot", width=2),
        )


def _build_base_kpi_figure(kpi, player_label, player_df_all):
    """Traccia di base (serie grezza del KPI lungo le 28 partite) condivisa
    da build_finding_chart e build_unified_finding_chart."""
    match_seq_full = list(range(N_REGULAR_MATCHES))
    sub = player_df_all[(player_df_all["player"] == player_label) & (player_df_all["kpi"] == kpi)]
    sub = sub.set_index("match_seq").reindex(match_seq_full)
    values = sub["value"].to_numpy()

    colors = theme_colors()
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=match_seq_full, y=values, mode="lines+markers",
        line=dict(color=colors["palette"][0], width=2), marker=dict(size=6),
        connectgaps=False, showlegend=False,
        customdata=[format_kpi_value(kpi, v) if not np.isnan(v) else "" for v in values],
        hovertemplate="%{x}<br>%{customdata}<extra></extra>",
    ))
    return fig, colors, match_seq_full


def build_finding_chart(kpi, player_label, finding, player_df_all, opponent_order):
    """
    Grafico di evidenza per un singolo finding: la serie grezza del KPI lungo
    tutta la stagione (28 partite), con la porzione che ha generato il
    finding evidenziata — verde = favorevole, rosso = sfavorevole (stessa
    convenzione status color di add_result_strip, qui applicata al
    finding invece che all'esito partita).
    """
    fig, colors, match_seq_full = _build_base_kpi_figure(kpi, player_label, player_df_all)
    _add_finding_annotation(fig, finding, colors)

    tick_labels = match_seq_tick_labels(opponent_order)
    fig.update_xaxes(tickmode="array", tickvals=match_seq_full, ticktext=tick_labels)
    fig.update_layout(**base_layout(f"{kpi} — {player_label}", "Avversario", kpi))
    fig.update_layout(height=320, margin=dict(t=60, b=40))
    return fig


def build_unified_finding_chart(kpi, player_label, findings, player_df_all, opponent_order):
    """
    Come build_finding_chart, ma con PIÙ finding sullo stesso KPI evidenziati
    su un unico grafico invece di uno per finding — usata dal report Word
    (richiesto esplicitamente dall'utente il 2026-08-15: "unifica gli
    eventi notevoli (streak o trend) in un unico grafico, se si riferiscono
    allo stesso kpi"), non dalla tab dashboard (che resta con un grafico per
    finding, già verificata).
    """
    fig, colors, match_seq_full = _build_base_kpi_figure(kpi, player_label, player_df_all)
    for finding in findings:
        _add_finding_annotation(fig, finding, colors)

    tick_labels = match_seq_tick_labels(opponent_order)
    fig.update_xaxes(tickmode="array", tickvals=match_seq_full, ticktext=tick_labels)
    fig.update_layout(**base_layout(f"{kpi} — {player_label}", "Avversario", kpi))
    fig.update_layout(height=320, margin=dict(t=60, b=40))
    return fig


EFFICIENCY_FAMILY_LABELS = {
    "ricezione": "Ricezione%",
    "attacco": "Attacco%",
    "attacco_so": "Attacco SO%",
    "attacco_fb": "Attacco FB%",
    "contrattacco": "Contrattacco%",
    "battuta": "Battuta%",
    "attacco_alzato_pos": ATTACCO_ALZATO_POS_KPI_LABELS["eff"],
    "attacco_alzato_escl": ATTACCO_ALZATO_ESCL_KPI_LABELS["eff"],
}


def render_season_summary(season_data, match_label_lookup):
    """
    Sintesi stagionale per un giocatore (src.player_season_report,
    richiesta esplicitamente dall'utente il 2026-08-14): efficienze
    aggregate e conteggi medi/partita curati per ruolo, punti di
    forza/debolezza (soglie assolute + confronto con i compagni), obiettivi
    per la prossima stagione (mediana di riferimento), partita migliore/
    peggiore per punteggio composito e per singolo KPI notevole.
    """
    st.subheader("Sintesi stagionale")
    st.caption(f"{season_data['n_partite_giocate']} partite giocate")

    efficienze = {k: v for k, v in season_data["efficienze"].items() if v[0] is not None}
    if efficienze:
        st.markdown("**Efficienze aggregate di stagione**")
        cols = st.columns(min(len(efficienze), 4))
        for i, (key, (eff, tot)) in enumerate(efficienze.items()):
            label = EFFICIENCY_FAMILY_LABELS.get(key, key)
            cols[i % len(cols)].metric(label, f"{eff:.1f}%", help=f"{tot} tentativi in stagione")

    conteggi = {k: v for k, v in season_data["conteggi_medi_partita"].items() if v[0] is not None}
    if conteggi:
        st.markdown("**Conteggi medi a partita**")
        table = pd.DataFrame([
            {"KPI": k, "Media/partita": round(media, 2), "Totale stagione": tot, "Partite con dati": n}
            for k, (media, tot, n) in conteggi.items()
        ])
        st.dataframe(table, width="stretch", hide_index=True)

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("**Punti di forza**")
        if season_data["punti_forza"]:
            for f in season_data["punti_forza"]:
                st.markdown(f"✅ **{f['area']}** — {f['motivo']}")
        else:
            st.caption("Nessuno di particolarmente notevole.")
    with col_b:
        st.markdown("**Punti deboli**")
        if season_data["punti_deboli"]:
            for f in season_data["punti_deboli"]:
                st.markdown(f"⚠️ **{f['area']}** — {f['motivo']}")
        else:
            st.caption("Nessuno di particolarmente notevole.")

    if season_data["obiettivi_prossima_stagione"]:
        st.markdown("**Obiettivi per la prossima stagione**")
        for t in season_data["obiettivi_prossima_stagione"]:
            label = EFFICIENCY_FAMILY_LABELS.get(t["kpi"], t["kpi"]) if t["tipo"] == "eff" else t["kpi"]
            unit = "%" if t["tipo"] == "eff" else ""
            st.markdown(f"🎯 {label}: {t['valore']:.1f}{unit} (mediana di riferimento {t['mediana']:.1f}{unit})")

    best, worst = season_data["partita_migliore"], season_data["partita_peggiore"]
    if best or worst:
        st.markdown(f"**Partita migliore/peggiore** ({season_data['criterio_migliore_peggiore']})")
        col_c, col_d = st.columns(2)
        if best:
            ref = match_label_lookup.get(best["match_seq"], f"partita {best['match_seq'] + 1}")
            if "netto" in best:
                col_c.metric("Migliore", f"+{best['netto']}", help=f"{ref} — {best['punti']} punti, {best['errori']} errori")
            else:
                col_c.metric("Migliore", f"{best['value']:.1f}%", help=ref)
        if worst:
            ref = match_label_lookup.get(worst["match_seq"], f"partita {worst['match_seq'] + 1}")
            if "netto" in worst:
                col_d.metric("Peggiore", f"{worst['netto']:+d}", help=f"{ref} — {worst['punti']} punti, {worst['errori']} errori")
            else:
                col_d.metric("Peggiore", f"{worst['value']:.1f}%", help=ref)

    if season_data["partite_notevoli"]:
        st.markdown("**Partite singole notevoli** (bilancio punti fatti − errori fatti, per fondamentale)")
        for n in season_data["partite_notevoli"]:
            ref = match_label_lookup.get(n["match_seq"], f"partita {n['match_seq'] + 1}")
            simbolo = "📈" if n["tipo"] == "migliore" else "📉"
            st.markdown(f"{simbolo} **Bilancio {n['kpi']}** — {n['tipo']} a {ref}: {n['value']:+.0f}")


# ----------------------------------------------------------------------------
# Report Word della pagella giocatore (pulsante "Genera report", richiesto
# esplicitamente dall'utente il 2026-08-15) — vedi generate_player_report_docx.
# ----------------------------------------------------------------------------

ROLE_DISPLAY_NAMES = {"L": "libero", "M": "schiacciatore", "C": "centrale", "O": "opposto", "P": "palleggiatore"}


def generate_narrative_summary(cognome, season_data, presenze):
    """
    Giudizio sintetico testuale in apertura del report Word — composto da
    regole a partire dalle statistiche già calcolate (punti di forza/
    debolezza, assenze, punteggio composito/ricezione, obiettivi), non un
    testo libero: ogni frase è tracciabile a un numero specifico del report.
    """
    ruolo_nome = ROLE_DISPLAY_NAMES.get(season_data["ruolo_bucket"], "giocatore")
    parti = [f"{cognome} ha giocato {season_data['n_partite_giocate']} partite nel ruolo di {ruolo_nome} nella stagione."]

    if presenze:
        tasso = presenze["tasso_assenza"] * 100
        if tasso == 0:
            parti.append("Presenza perfetta ad allenamenti e partite, senza assenze registrate.")
        elif tasso < 10:
            parti.append(f"Presenza regolare, con un tasso di assenza contenuto ({tasso:.1f}%).")
        elif tasso < 20:
            parti.append(f"Tasso di assenza nella norma ({tasso:.1f}%).")
        else:
            parti.append(f"Tasso di assenza elevato nel corso della stagione ({tasso:.1f}%), un aspetto da monitorare.")

    punti_forza = season_data["punti_forza"]
    punti_deboli = season_data["punti_deboli"]
    if punti_forza and punti_deboli:
        parti.append(
            f"Sul piano del rendimento, i punti di forza principali sono {', '.join(f['area'] for f in punti_forza)}, "
            f"mentre le aree su cui lavorare restano {', '.join(f['area'] for f in punti_deboli)}."
        )
    elif punti_forza:
        parti.append(
            f"Sul piano del rendimento emergono soprattutto i punti di forza in "
            f"{', '.join(f['area'] for f in punti_forza)}, senza criticità particolarmente marcate rispetto ai compagni di ruolo."
        )
    elif punti_deboli:
        parti.append(
            f"Sul piano del rendimento le aree principali su cui lavorare sono "
            f"{', '.join(f['area'] for f in punti_deboli)}, senza punti di forza particolarmente marcati rispetto ai compagni di ruolo."
        )
    else:
        parti.append("Il rendimento stagionale resta nella norma, senza punti di forza o criticità particolarmente marcati rispetto ai compagni di ruolo.")

    best, worst = season_data["partita_migliore"], season_data["partita_peggiore"]
    if best and worst:
        criterio = season_data["criterio_migliore_peggiore"]
        if "netto" in best:
            parti.append(
                f"La partita migliore della stagione (per {criterio.lower()}) ha visto un bilancio di "
                f"{best['netto']:+d}, contro il {worst['netto']:+d} della partita peggiore."
            )
        else:
            parti.append(
                f"La partita migliore della stagione (per {criterio.lower()}) ha toccato il {best['value']:.1f}%, "
                f"contro il {worst['value']:.1f}% della partita peggiore."
            )

    n_obiettivi = len(season_data["obiettivi_prossima_stagione"])
    if n_obiettivi:
        parti.append(f"Per la prossima stagione emergono {n_obiettivi} obiettivi specifici di miglioramento, dettagliati più sotto.")

    return " ".join(parti)


def _docx_add_markdown_bold_paragraph(doc, text):
    """
    Aggiunge un paragrafo interpretando i marcatori '**bold**' in stile
    Markdown (usati da format_finding_text, pensato per st.markdown) come
    grassetto reale — altrimenti nel documento Word restano asterischi
    letterali (bug riscontrato in verifica: "**Attacco SO%** — periodo...").
    """
    p = doc.add_paragraph()
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = bool(i % 2)
    return p


def _docx_add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = str(header)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = str(value)
    return table


def generate_player_report_docx(cognome, data, season_data, player_df_all, team_df, opponent_order, match_label_lookup):
    """
    Assembla il report Word della pagella giocatore, sezione per sezione
    (ordine richiesto esplicitamente dall'utente il 2026-08-15): giudizio
    sintetico, ruolo/assenze, sintesi stagionale, conteggi medi a partita,
    punti di forza/deboli (sempre presenti, anche vuoti), partita migliore/
    peggiore e partite notevoli, grafici dell'andamento (un grafico per KPI,
    non per finding — vedi build_unified_finding_chart), obiettivi per la
    prossima stagione (senza la spiegazione della mediana, a differenza
    della tab dashboard). Ritorna i byte del documento .docx.
    """
    doc = Document()
    doc.add_heading(f"Pagella — {cognome}", level=0)

    doc.add_heading("Giudizio sintetico", level=1)
    doc.add_paragraph(generate_narrative_summary(cognome, season_data, data["presenze"]))

    doc.add_heading("Ruolo e presenze", level=1)
    pres = data["presenze"]
    doc.add_paragraph(f"Ruolo: {data['ruolo'] or 'n/d'}")
    if pres:
        doc.add_paragraph(
            f"Assenza: {pres['tasso_assenza'] * 100:.1f}% — presente/convocato: {pres['n_presente']}/{pres['n_convocato']}"
        )
    else:
        doc.add_paragraph("Presenze: nessun dato")

    doc.add_heading("Sintesi stagionale", level=1)
    doc.add_paragraph(f"{season_data['n_partite_giocate']} partite giocate")
    efficienze = {k: v for k, v in season_data["efficienze"].items() if v[0] is not None}
    if efficienze:
        _docx_add_table(
            doc, ["KPI", "Efficienza", "Tentativi in stagione"],
            [(EFFICIENCY_FAMILY_LABELS.get(k, k), f"{eff:.1f}%", tot) for k, (eff, tot) in efficienze.items()],
        )

    doc.add_heading("Conteggi medi a partita", level=1)
    conteggi = {k: v for k, v in season_data["conteggi_medi_partita"].items() if v[0] is not None}
    if conteggi:
        _docx_add_table(
            doc, ["KPI", "Media/partita", "Totale stagione", "Partite con dati"],
            [(k, f"{media:.2f}", tot, n) for k, (media, tot, n) in conteggi.items()],
        )
    else:
        doc.add_paragraph("Nessun conteggio disponibile per questo ruolo.")

    doc.add_heading("Punti di forza", level=1)
    if season_data["punti_forza"]:
        for f in season_data["punti_forza"]:
            doc.add_paragraph(f"{f['area']} — {f['motivo']}", style="List Bullet")
    else:
        doc.add_paragraph("Nessuno di particolarmente notevole.")

    doc.add_heading("Punti deboli", level=1)
    if season_data["punti_deboli"]:
        for f in season_data["punti_deboli"]:
            doc.add_paragraph(f"{f['area']} — {f['motivo']}", style="List Bullet")
    else:
        doc.add_paragraph("Nessuno di particolarmente notevole.")

    doc.add_heading("Partita migliore e peggiore", level=1)
    best, worst = season_data["partita_migliore"], season_data["partita_peggiore"]
    if best or worst:
        doc.add_paragraph(f"Criterio: {season_data['criterio_migliore_peggiore']}")
        if best:
            ref = match_label_lookup.get(best["match_seq"], f"partita {best['match_seq'] + 1}")
            if "netto" in best:
                doc.add_paragraph(f"Migliore: {ref} — {best['punti']} punti, {best['errori']} errori (netto {best['netto']:+d})", style="List Bullet")
            else:
                doc.add_paragraph(f"Migliore: {ref} — {best['value']:.1f}%", style="List Bullet")
        if worst:
            ref = match_label_lookup.get(worst["match_seq"], f"partita {worst['match_seq'] + 1}")
            if "netto" in worst:
                doc.add_paragraph(f"Peggiore: {ref} — {worst['punti']} punti, {worst['errori']} errori (netto {worst['netto']:+d})", style="List Bullet")
            else:
                doc.add_paragraph(f"Peggiore: {ref} — {worst['value']:.1f}%", style="List Bullet")
    else:
        doc.add_paragraph("Nessun dato sufficiente.")

    doc.add_heading("Partite notevoli", level=1)
    if season_data["partite_notevoli"]:
        for n in season_data["partite_notevoli"]:
            ref = match_label_lookup.get(n["match_seq"], f"partita {n['match_seq'] + 1}")
            doc.add_paragraph(f"Bilancio {n['kpi']} — {n['tipo']} a {ref}: {n['value']:+.0f}", style="List Bullet")
    else:
        doc.add_paragraph("Nessuna di particolarmente notevole.")

    doc.add_heading("Andamento nel corso della stagione", level=1)
    player_label = data["player_label"]
    findings = data["findings"]
    if not findings or player_label is None:
        doc.add_paragraph("Nessun finding notevole per questo giocatore.")
    else:
        findings_by_kpi = {}
        for finding in findings:
            findings_by_kpi.setdefault(finding["kpi"], []).append(finding)
        for kpi, kpi_findings in findings_by_kpi.items():
            doc.add_heading(kpi, level=2)
            for finding in kpi_findings:
                _docx_add_markdown_bold_paragraph(doc, format_finding_text(finding, match_label_lookup))
            fig = build_unified_finding_chart(kpi, player_label, kpi_findings, player_df_all, opponent_order)
            png_bytes = fig.to_image(format="png", width=900, height=320, scale=2)
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(6))

    doc.add_heading("Obiettivi per la prossima stagione", level=1)
    if season_data["obiettivi_prossima_stagione"]:
        for t in season_data["obiettivi_prossima_stagione"]:
            label = EFFICIENCY_FAMILY_LABELS.get(t["kpi"], t["kpi"]) if t["tipo"] == "eff" else t["kpi"]
            unit = "%" if t["tipo"] == "eff" else ""
            doc.add_paragraph(f"{label}: {t['valore']:.1f}{unit}", style="List Bullet")
    else:
        doc.add_paragraph("Nessun obiettivo specifico individuato.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_player_report_section(report, not_in_registry, season_report, player_df_all, team_df, opponent_order):
    st.header("Pagella giocatore")
    st.caption(
        "Presenze/assenze e finding di rendimento più notevoli per giocatore — streak (media mobile "
        "di 3 partite giocate consecutive sopra/sotto la mediana stagionale, per intercettare anche "
        "una crescita/calo momentaneo) o cambio di livello (il punto di rottura più marcato della "
        "stagione, non necessariamente a metà), su un sottoinsieme curato di KPI (percentuali "
        "principali + pochi conteggi mirati). Per i KPI percentuali, le partite con troppi pochi "
        "tentativi (< 5 su circa 120 azioni totali a partita) sono escluse dal calcolo. Per i "
        "palleggiatori la famiglia attacco generica è sostituita da \"Attacco Alzato\": l'efficienza "
        "di attacco dei loro attaccanti nelle sole azioni in cui il palleggiatore in campo è "
        "identificabile con sufficiente sicurezza."
    )
    if not_in_registry:
        st.caption(
            "Nel foglio presenze ma non nel registro giocatori (non mostrati qui): "
            + ", ".join(not_in_registry)
        )

    cognomi = sorted(report.keys())
    if not cognomi:
        st.info("Nessun giocatore trovato.")
        return

    match_label_lookup = build_match_label_lookup(team_df)
    tabs = st.tabs(cognomi)
    for tab, cognome in zip(tabs, cognomi):
        with tab:
            data = report[cognome]
            pres = data["presenze"]

            col1, col2, col3 = st.columns(3)
            col1.metric("Ruolo", data["ruolo"] or "n/d")
            if pres:
                col2.metric("Assenza", f"{pres['tasso_assenza'] * 100:.1f}%")
                col3.metric("Presente / Convocato", f"{pres['n_presente']}/{pres['n_convocato']}")
            else:
                col2.metric("Assenza", "n/d")

            season_data = season_report.get(cognome)
            if season_data:
                st.divider()
                render_season_summary(season_data, match_label_lookup)

            st.divider()
            st.subheader("Andamento nel corso della stagione")
            player_label = data["player_label"]
            if not data["findings"]:
                st.info("Nessun finding notevole per questo giocatore (dati insufficienti o troppo stabili).")
            elif player_label is None:
                st.warning("Nessuna serie KPI disponibile per questo giocatore.")
            else:
                for finding in data["findings"]:
                    st.markdown(format_finding_text(finding, match_label_lookup))
                    fig = build_finding_chart(finding["kpi"], player_label, finding, player_df_all, opponent_order)
                    st.plotly_chart(fig, theme=None)

            if season_data:
                st.divider()
                report_key = f"report_docx_{cognome}"
                if st.button("📄 Genera report", key=f"genera_report_{cognome}"):
                    with st.spinner("Generazione report Word..."):
                        st.session_state[report_key] = generate_player_report_docx(
                            cognome, data, season_data, player_df_all, team_df, opponent_order, match_label_lookup,
                        )
                if report_key in st.session_state:
                    st.download_button(
                        "⬇️ Scarica report",
                        data=st.session_state[report_key],
                        file_name=f"pagella_{cognome}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"scarica_report_{cognome}",
                    )


def main():
    st.title("Decimo Roma — Confronto Andata vs Ritorno")
    st.caption(f"Stagione {SEASON} · playout (solo JVC) mostrato come 14ª posizione / 27ª-28ª partita")

    with st.sidebar:
        st.header("Impostazioni generali")
        st.slider(
            "Smoothing linee di trend (kernel gaussiano)", 0.0, 3.0, 1.2, 0.1,
            key="bandwidth",
            help="0 = linea grezza punto a punto; valori più alti = curva più smussata",
        )
        selected_rec_vote = st.multiselect(
            'Esito ricezione per "Attacco SO"', REC_VOTE_OPTIONS,
            default=list(DEFAULT_REC_VOTE), key="rec_vote",
            help=(
                'Un attacco è "SO" (side-out) se segue immediatamente una ricezione con uno '
                'di questi voti (al più con un\'alzata di mezzo) — vedi i KPI "Attacco SO...". '
                "Cambiare la selezione ricalcola quei 4 KPI."
            ),
        )
        if not selected_rec_vote:
            st.warning('Nessun esito ricezione selezionato: i KPI "Attacco SO..." saranno vuoti.')

        min_attacks = st.number_input(
            "Soglia minima attacchi (per giocatore/partita)", min_value=0, value=0, step=1,
            help=(
                'Per i KPI della famiglia "attacco" (Attacco*, Attacco SO*, Contrattacco*): se in una '
                "partita un giocatore non raggiunge questo numero di attacchi della stessa famiglia, "
                "per quella partita risulta senza dati per quel KPI (come se non avesse giocato quel "
                "fondamentale) — non uno zero. Non si applica a Battuta/Ricezione/Muro/Errori. 0 = nessun filtro."
            ),
        )
        st.divider()

    team_df_fixed, player_df_fixed = load_data(SEASON)
    team_df_so, player_df_so = load_attacco_so_data(SEASON, tuple(selected_rec_vote))
    team_df = pd.concat([team_df_fixed, team_df_so], ignore_index=True)
    setter_df = load_setter_kpi_data(SEASON)
    # player_df_all: senza soglia minima attacchi — è la base su cui src.player_report
    # ha calcolato i finding della pagella, i grafici di evidenza devono usare la
    # stessa serie (altrimenti un finding potrebbe riferirsi a partite che nel
    # grafico risultano "senza dati" per effetto della soglia della sidebar). Include
    # anche "Attacco Alzato" (ruolo palleggiatore): non soggetto alla soglia minima
    # attacchi (quella si applica solo alle famiglie di attacco generiche).
    player_df_all = pd.concat([player_df_fixed, player_df_so, setter_df], ignore_index=True)
    player_df = apply_min_attacks_threshold(player_df_all, min_attacks)

    match_outcomes_df = load_match_outcomes(SEASON)
    rec_vote_for_report = tuple(selected_rec_vote) if selected_rec_vote else DEFAULT_REC_VOTE
    player_report, player_report_missing = load_player_report(SEASON, rec_vote_for_report)
    season_report = load_player_season_report(SEASON, rec_vote_for_report)

    x_axis_order = get_x_axis_order(SEASON)
    players = sorted(player_df["player"].unique())
    entity_options = ["Squadra"] + players
    opponent_order = x_axis_order[:-1]  # senza PLAYOUT_LABEL, per il grafico 2

    render_kpi_section(team_df, player_df, x_axis_order, entity_options)
    st.divider()
    render_comparison_section(team_df, player_df, entity_options, opponent_order, match_outcomes_df)
    st.divider()
    render_kpi_comparison_section(team_df, player_df, entity_options, opponent_order, match_outcomes_df)
    st.divider()
    render_player_report_section(
        player_report, player_report_missing, season_report, player_df_all, team_df, opponent_order,
    )


if __name__ == "__main__":
    main()
